from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

# Profile Picture
document.add_picture('assets/me.jpg', width=Inches(2.0))

# Personal Information Input
speak('Welcome to My CV Builder!')
speak('Please enter your name.')
name = input('Enter your name: ')

speak(f'Hello, {name}! Let\'s build your CV together.')
speak('Please enter your phone number.')
phone_number = input('Enter your phone number: ')

speak('Please enter your email address.')
email = input('Enter your email: ')

# Add Personal Information
document.add_paragraph(
  name + ' | ' + phone_number + ' | ' + email
)

# About Me Section
document.add_heading('About Me')
# about_me = input('Tell me about yourself: ')
speak(f'Tell me about yourself, {name}.')
document.add_paragraph(input('Tell me about yourself: '))

# Work Experience Section
speak(f'Let\'s add your work experience, {name}.')
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Enter company name.')
company = input('Enter company name: ')

speak('From Date.')
from_date = input('From Date: ')

speak('To Date.')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

speak('Describe your experience at ' + company + '.')
experience_details = input('Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# More experiences
while True:
    speak('Do you have more work experience to add? Yes or no.')
    more_experience = input(
        'Do you have more work experience to add? Yes or no: '
    )
    if more_experience.lower() == 'yes':
        document.add_heading('Work Experience')
        p = document.add_paragraph()

        speak('Enter company name.')
        company = input('Enter company name: ')

        speak('From Date.')
        from_date = input('From Date: ')

        speak('To Date.')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        speak('Describe your experience at ' + company + '.')
        experience_details = input('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break

# Skills Section

document.add_heading('Skills')
speak('Enter your skills. Please separate by commas.')
skills = input('Enter your skills. Please separate by commas: ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

# More skills
while True:
    speak('Do you have more skills to add? Yes or no.')
    more_skills = input('Do you have more skills to add? Yes or no: ')
    if more_skills.lower() == 'yes':
        speak('Enter your skills. Please separate by commas.')
        skills = input('Enter your skills. Please separate by commas: ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        speak('Thank you for using My CV Builder. Your CV is being generated.')
        speak('Goodbye!')
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using My CV Builder"

document.save('cv.docx')